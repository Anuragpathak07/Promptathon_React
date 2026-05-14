import os
import sys

def generate_docx():
    try:
        import docx
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls, qn
    except ImportError:
        print("python-docx package not found. Installing now...")
        os.system("pip install python-docx")
        import docx
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls, qn

    doc = Document()

    # Page Setup - Normal Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette - Executive Theme
    COLOR_PRIMARY = RGBColor(14, 22, 40)     # Deep Navy #0E1628
    COLOR_SECONDARY = RGBColor(74, 191, 168) # Teal #4ABFA8
    COLOR_ACCENT = RGBColor(232, 103, 74)    # Coral #E8674A
    COLOR_TEXT = RGBColor(40, 40, 40)        # Charcoal #282828
    COLOR_MUTED = RGBColor(120, 120, 120)    # Muted Grey #787878

    # Style Configurations
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = COLOR_TEXT

    # Helper function for headings
    def add_heading_styled(text, level):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        
        if level == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = COLOR_PRIMARY
            # Add bottom border under Heading 1
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                             r'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="4ABFA8"/>'
                             r'</w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
        elif level == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = COLOR_SECONDARY
        elif level == 3:
            run.font.size = Pt(12.5)
            run.font.color.rgb = COLOR_ACCENT
        return p

    # Title Page / Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("Enterprise AI Quality Control Platform")
    run_title.bold = True
    run_title.font.size = Pt(28)
    run_title.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(36)
    run_sub = p_sub.add_run("End-to-End Technical Architecture, Mathematical Foundations, and Production Engineering Report")
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = COLOR_MUTED

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(48)
    run_meta = p_meta.add_run("Author: Senior AI Systems Architect & ML Lead\nDocument Classification: Enterprise Architecture & Research Thesis\nSystem Version: Production v4.2")
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = COLOR_MUTED

    doc.add_page_break()

    # Content Data
    content = [
        ("Executive Summary & System Architecture Overview", 1,
         "The Industrial Quality Control Platform is a state-of-the-art hybrid artificial intelligence vision framework designed for real-time anomaly detection, zero-shot defect classification, and automated engineering specification compliance. The architecture bridges the gap between deterministic deep feature representations (PatchCore) and semantic reasoning engines (Vision Large Language Models).\n\n"
         "Operating in high-throughput industrial manufacturing environments—such as Surface Mount Technology (SMT) printed circuit board assembly lines, precision mechanical fabrication, and pharmaceutical packaging—the platform provides sub-second defect detection with zero positive baseline drift."),
        
        ("Module 1: Ingestion & Multi-modal Data Engineering Pipeline", 1, ""),
        ("Purpose", 2, "The ingestion subsystem acts as the deterministic data gateway for the anomaly detection platform. It is responsible for fetching, validating, structuring, and pre-processing industrial image datasets from cloud repositories (HuggingFace Hub via Voxel51) and local directories. It ensures that downstream feature extractors receive perfectly standardized tensors invariant to camera resolution variances."),
        ("Internal Workflow & Algorithms", 2, 
         "1. Cache Redirection: To prevent operating system permission blocks, environment variables redirect all data streaming into a local writable directory (./hf_cache).\n"
         "2. Deterministic Splitting: When reading annotation files (image_anno.csv), the module separates pristine ('normal') samples from anomalous samples. The normal samples are deterministically split into an 80% training coreset and a 20% validation split.\n"
         "3. Anti-Aliased Downsampling: Source images are loaded via PIL and resized to 224x224 using the Lanczos-3 filter (sinc-based resampling), which preserves high-frequency edge features crucial for defect detection.\n"
         "4. Standardization: Tensors are normalized using ImageNet channel-wise statistics (mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225])."),
        ("Architecture Rationale & Tradeoffs", 2, 
         "Lanczos vs. Bilinear/Bicubic: Bilinear interpolation causes spatial blurring, which can mask micro-scratches on metallic traces. Lanczos retains sharpness but incurs a ~15% higher CPU computational cost during preprocessing.\n\n"
         "Offline Caching: Streaming directly from HuggingFace during training introduces network latency and risk of connection drops. Disk-caching raw PNGs guarantees 100% reproducible training runs at maximum disk read throughput."),
        ("Runtime Behavior & Production Bottlenecks", 2, 
         "Time Complexity: O(H x W) per image for downsampling and channel normalization.\n"
         "Memory Usage: ~600 MB RAM overhead during dataset indexing and CSV loading.\n"
         "Memory Hangs on Windows: Multiprocessing DataLoader with num_workers > 0 on Windows platforms causes spawn deadlocks due to Python GIL constraints. The module intercepts Windows platforms and forces num_workers = 0 to ensure stable single-threaded throughput."),

        ("Module 2: AI Semantic Segmentation & Caching Subsystem", 1, ""),
        ("Purpose", 2, "Industrial vision stations suffer from unconstrained background variations, such as scratched conveyor belts, green cutting mats, operator hands, or shadow fluctuations. If passed directly to feature extractors, the neural network constructs feature embeddings of the background table, causing catastrophic false-positive rates. This module performs deep semantic background subtraction while preserving delicate foreground geometry (e.g., metallic PCB pins)."),
        ("Internal Workflow & Algorithms", 2, 
         "1. MD5 Hashing Cache: To eliminate redundant segmentation, source image paths are hashed via MD5. If a cached PNG exists in ./hf_cache/rembg_cache, it is loaded instantly.\n"
         "2. Aggressive Downscaling: High-resolution source images (4000x3000) are downscaled to 160x160 before inference.\n"
         "3. U2-Net Deep Inference: The downscaled image is passed to rembg running an ONNX runtime session of u2netp (a highly pruned 4MB variant of U2-Net).\n"
         "4. Alpha Compositing: The resulting RGBA mask is composited onto a solid black RGB canvas ([0, 0, 0]), isolating the pristine component.\n"
         "5. Upsampling & Caching: The black-background image is upscaled back to original resolution and saved to disk."),
        ("Architecture Rationale & Tradeoffs", 2, 
         "Pruned u2netp vs. Standard u2net: Standard U2-Net is a 176MB model requiring ~3 seconds per image on CPU. u2netp is 4MB and executes in ~15 milliseconds while maintaining identical structural boundary accuracy for rigid industrial components.\n\n"
         "AI Masking vs. Grayscale Otsu: Otsu thresholding relies strictly on pixel brightness. Highly reflective metallic pins share the same brightness as a grey table and get truncated by Otsu contour extraction. U2-Net understands object semantics and preserves 100% of metallic pins."),

        ("Module 3: Deep Feature Extraction Layer", 1, ""),
        ("Purpose", 2, "To detect anomalies without prior knowledge of defect types, the system requires a rich, multi-scale feature representation of 'normal' geometry. This module utilizes a deep convolutional network to project spatial image regions into highly discriminative embedding vectors."),
        ("Internal Workflow & Algorithms", 2, 
         "1. Backbone Architecture: ResNet-50 pre-trained on ImageNet-1k is deployed as the feature backbone. The classification head (fc) and global pooling layers are severed.\n"
         "2. Forward Feature Hooks: PyTorch forward hooks are registered to intermediate residual stages: layer2 (downsampling stride 8) and layer3 (downsampling stride 16).\n"
         "3. Multi-Scale Feature Interpolation: Output maps from layer3 (shape B, 1024, 14, 14) are bilinearly upsampled to match the spatial dimensions of layer2 (shape B, 512, 28, 28).\n"
         "4. Concatenation & Normalization: Channel representations are concatenated to form a (B, 1536, 28, 28) tensor, which is then L2-normalized across the channel dimension."),
        ("Architecture Rationale & Tradeoffs", 2, 
         "ResNet-50 vs. Vision Transformers (ViT): ViT patch embeddings lack translation invariance and require massive computational budgets at high resolutions. ResNet-50 provides highly robust spatial inductive biases at a fraction of the memory footprint.\n\n"
         "Layer Selection (layer2 + layer3): layer1 captures low-level edge noise (e.g., natural surface texture), leading to false positives. layer4 captures abstract semantic concepts (e.g., 'objectness') but loses exact spatial localization. layer2 and layer3 represent the optimal balance between spatial resolution and structural semantics."),

        ("Module 4: Coreset Sub-sampling & FAISS Memory Bank Architecture", 1, ""),
        ("Purpose", 2, "Extracting patch features across thousands of training images generates millions of 1536-dimensional vectors. Storing and searching this raw matrix in real-time is computationally intractable. This module constructs a compressed 'Coreset' memory bank that preserves the geometric topology of the feature space while discarding 95% of redundant embeddings."),
        ("Internal Workflow & Algorithms", 2, 
         "1. Local Neighbourhood Aggregation: To increase spatial robustness against minor misalignment, intermediate feature maps undergo adaptive average pooling with kernel size 3x3 and stride 1.\n"
         "2. Dimensionality Reshaping: Spatial tensors are flattened into a dense feature matrix M of size (N_images x 28 x 28, 1536).\n"
         "3. Coreset Subsampling: The system implements greedy approximate k-centre coverage (or fast uniform random sampling at ratio 0.05) to extract a representative coreset C.\n"
         "4. FAISS Indexing: The coreset is indexed inside Facebook AI Similarity Search (faiss.IndexFlatL2) for rapid Euclidean nearest-neighbour distance calculations.\n"
         "5. Threshold Calibration: The module evaluates the training set against the constructed index, extracting the 99th percentile of patch distances as the global anomaly threshold."),
        ("Mathematical Foundations", 2, 
         "The anomaly score s* for a test image x with patch features P(x) against coreset C is calculated as the maximum nearest-neighbour distance:\n"
         "s* = max_{p in P(x)} min_{c in C} || p - c ||_2\n\n"
         "To ensure robustness against single-patch outlier noise, the system modifies the raw maximum to the 99th percentile distance across all test patches."),

        ("Module 5: Spatial Anomaly Heatmap & Distance Scoring Pipeline", 1, ""),
        ("Purpose", 2, "Detecting an anomaly is insufficient for industrial operations; engineers require exact spatial localization of the defect. This module transforms discrete patch distance scores into a smooth, high-resolution anomaly heatmap overlay."),
        ("Internal Workflow & Algorithms", 2, 
         "1. Spatial Unflattening: The 784 nearest-neighbour distance scores extracted from FAISS are reshaped back into a 28x28 spatial grid.\n"
         "2. Bilinear Upsampling: The low-resolution distance map is bilinearly upsampled to match the source image resolution (224x224 or 1024x1024).\n"
         "3. Min-Max Normalization: The upsampled map is normalized against the global calibrated threshold.\n"
         "4. Colormap Blending: The normalized heatmap is projected into the JET colormap space (blue = cool/normal, red = warm/anomalous) and alpha-blended with the original RGB tensor at alpha = 0.5."),

        ("Module 6: Agentic Vision LLM Inspection Copilot", 1, ""),
        ("Purpose", 2, "While PatchCore provides mathematical anomaly detection, it cannot explain why an anomaly occurred or what specifications were violated. This module integrates Vision Large Language Models (Groq Llama-3-Vision / OpenAI GPT-4o) as an expert industrial engineering copilot to generate structured inspection reports."),
        ("Internal Workflow & Algorithms", 2, 
         "1. Base64 Encoding: The masked PIL image is compressed and encoded into a Base64 JPEG string.\n"
         "2. Contextual Prompt Engineering: The module constructs an advanced prompt injecting the exact MVTec component category name, PatchCore anomaly score, and numerical threshold.\n"
         "3. Strict Formatting Guardrails: The system prompt explicitly forbids Markdown syntax to ensure clean ingestion by custom HTML parsers.\n"
         "4. Structured Output Extraction: The AI model generates an inspection report divided into deterministic numbered sections: 1. CURRENT STATUS, 2. SPEC COMPLIANCE, 3. RISK LEVEL, 4. RECOMMENDATIONS."),

        ("Module 7: Zero-Shot Fallback AI Subsystem", 1, ""),
        ("Purpose", 2, "When an operator uploads an image of an unconfigured or brand-new component category (where no PatchCore .pkl model exists), the system cannot fail. This subsystem intercepts untrained categories and routes them to a zero-shot Vision AI pipeline."),
        ("Internal Workflow", 2, 
         "1. JSON Schema Enforcement: The LMM is instructed to output a single valid JSON object containing exact keys: is_anomaly, anomaly_score, verdict_reason, and defects.\n"
         "2. Normalized Bounding Boxes: Detected defects are returned with bounding box coordinates box_2d in normalized percentage values [ymin, xmin, ymax, xmax] from 0 to 100.\n"
         "3. Synthetic Map Generation: To ensure seamless UI compatibility, the module synthesizes a 28x28 spatial heatmap, illuminating the exact bounding box regions with high anomaly scores (0.85+)."),

        ("Module 8: Enterprise Web UI & Rich Analytics Dashboard", 1, ""),
        ("Purpose", 2, "The user interface serves as the mission-control dashboard for factory operators and quality engineers. Built on Gradio, it renders high-fidelity visual overlays, real-time statistical breakdowns, and historical detection trends."),
        ("Internal Workflow & Visual Architecture", 2, 
         "1. Animated SVG Score Ring: Computes arc circumferences dynamically to render a circular progress gauge illustrating the anomaly score against the threshold.\n"
         "2. Spatial Zone Analysis: Divides the spatial heatmap into a 3x3 matrix, computing mean anomaly scores per physical quadrant to instantly identify localized manufacturing drift.\n"
         "3. Distribution Histogram: Flattens the 28x28 score map into 12 distinct density buckets, plotting an inline SVG histogram with a dynamic threshold marker.\n"
         "4. Historical Sparkline: Retains an in-memory queue of the last 20 detections, rendering a real-time SVG sparkline tracking anomaly score drift over time."),

        ("System Deployment, Security, Bottlenecks & Future Roadmap", 1, ""),
        ("Enterprise Deployment Concerns", 2, "Air-Gapped Factory Floors: Cloud API dependencies (Groq/OpenAI) fail in strictly air-gapped industrial facilities. Deployment requires swapping cloud endpoints with local TensorRT-LLM or vLLM instances running quantized Llama-3-Vision models on edge Industrial PCs."),
        ("Production Bottlenecks", 2, "GIL & CPU Concurrency: Python's Global Interpreter Lock restricts multi-core scaling for CPU-bound DataLoader processing. Enterprise scaling requires migrating core inference pipelines to C++ using LibTorch and ONNX Runtime C APIs."),
        ("Security Implications", 2, "Model Poisoning: If an adversary compromises the training image directory, injecting subtle defects into the 'normal' folder, PatchCore will incorporate those defects into its coreset memory bank, permanently blinding the system to those flaws. Mitigation requires cryptographic hashing and strict access control on all training data repositories."),
        ("Future Improvements", 2, 
         "1. Active RAG Integration: Connect Vision LLMs to vector databases containing proprietary manufacturer CAD schematics, bill of materials (BOM), and IPC-A-610 standards for highly rigorous compliance verification.\n"
         "2. Autonomous MES Ticketing: Implement LangChain/CrewAI agentic wrappers to automatically trigger SAP / ServiceNow maintenance tickets when critical risk thresholds are breached.")
    ]

    for title, level, body in content:
        add_heading_styled(title, level)
        if body:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.15
            p.add_run(body)

    output_path = os.path.abspath("Enterprise_AI_Quality_Control_Architecture_Report.docx")
    doc.save(output_path)
    print(f"\nSUCCESS: Enterprise Word Document generated at:\n{output_path}\n")

if __name__ == "__main__":
    generate_docx()
